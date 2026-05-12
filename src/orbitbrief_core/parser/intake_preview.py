from __future__ import annotations

from dataclasses import replace
import re
from pathlib import Path
from typing import Any, Mapping

from orbitbrief_core.parser.adapters.spreadsheet_common import build_spreadsheet_preview
from orbitbrief_core.parser.router import RouterInput

_MAX_PREVIEW_CHARS = 12000
_TEMPLATE_PLACEHOLDER_RE = re.compile(r"\{\{[^}]+\}\}")

_META_REFERENCE_SIGNAL_TOKENS: tuple[str, ...] = (
    "orbitbrief",
    "parser architecture",
    "evidence operating system",
    "multimodal evidence compiler",
    "bounded claim extraction",
    "bounded narrative extraction",
    "compiled contract policy",
    "compiled runtime policy",
    "graph-backed packet",
    "packetizer",
    "qwen",
    "layer 1",
    "layer 2",
    "layer 3",
    "layer 4",
    "layer 5",
    "layer 6",
    "layer 7",
    "layer 8",
    "layer 9",
    "layer 10",
    "layer 11",
    "modality-native recovery",
    "projection + deterministic postprocess",
)


def hydrate_router_input(router_input: RouterInput) -> RouterInput:
    metadata = dict(router_input.metadata) if isinstance(router_input.metadata, Mapping) else {}
    raw_text = str(metadata.get("raw_text") or "").strip()
    preview = str(router_input.raw_text_preview or "").strip() or raw_text
    path = _resolve_path(router_input)

    if path is not None:
        suffix = path.suffix.lower()
        if not preview:
            if suffix in {".txt", ".md", ".eml"}:
                preview = path.read_text(encoding="utf-8", errors="replace")
                metadata.setdefault("raw_text", preview)
            elif suffix == ".docx":
                full_text = _docx_preview_text(path)
                if full_text:
                    preview = full_text[:_MAX_PREVIEW_CHARS]
                    metadata.setdefault("raw_text", full_text)
                    metadata.setdefault("full_text", full_text)
            elif suffix in {".xlsx", ".csv"}:
                spreadsheet_preview = build_spreadsheet_preview(path, max_preview_chars=_MAX_PREVIEW_CHARS)
                preview = spreadsheet_preview.preview_text
                if spreadsheet_preview.full_text:
                    metadata.setdefault("raw_text", spreadsheet_preview.full_text)
                    metadata.setdefault("full_text", spreadsheet_preview.full_text)
                metadata.setdefault("spreadsheet_relevant_sheets", list(spreadsheet_preview.relevant_sheet_names))
                metadata.setdefault("spreadsheet_skipped_sheets", list(spreadsheet_preview.skipped_sheet_names))
                metadata.setdefault("spreadsheet_block_count", spreadsheet_preview.block_count)
            elif suffix == ".pdf":
                use_full_pdf = bool(metadata.get("cad_hint")) or bool(metadata.get("drawing_packet"))
                full_text = _pdf_preview_text(path, full_document=use_full_pdf)
                if full_text:
                    preview = full_text[:_MAX_PREVIEW_CHARS]
                    metadata.setdefault("raw_text", full_text)
                    metadata.setdefault("full_text", full_text)
                    token_count = len([token for token in full_text.split() if token.strip()])
                    metadata.setdefault("native_text_ratio", 1.0 if token_count >= 40 else 0.1)
        else:
            metadata.setdefault("raw_text", preview)

    template_meta = _detect_template_schema_artifact(preview or raw_text)
    if template_meta:
        metadata.update(template_meta)

    meta_reference = _detect_meta_reference_artifact(preview or raw_text)
    if meta_reference:
        metadata.update(meta_reference)

    if not preview:
        preview = str(metadata.get("raw_text") or "")[:_MAX_PREVIEW_CHARS]

    return replace(router_input, raw_text_preview=preview, metadata=metadata)


def _resolve_path(router_input: RouterInput) -> Path | None:
    if isinstance(router_input.metadata, Mapping):
        for key in ("path", "file_path", "local_path", "source_path"):
            value = router_input.metadata.get(key)
            if value:
                path = Path(str(value))
                if path.exists():
                    return path
    if router_input.filename:
        path = Path(str(router_input.filename))
        if path.exists():
            return path
    return None


def _docx_preview_text(path: Path) -> str:
    from xml.etree import ElementTree as ET
    from zipfile import ZipFile

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    def _local_name(tag: str) -> str:
        return tag.split("}", 1)[1] if "}" in tag else tag

    def _xml_text(element: ET.Element) -> str:
        parts: list[str] = []
        for node in element.iter():
            if _local_name(node.tag) == "t" and node.text:
                parts.append(node.text)
        return " ".join(" ".join(parts).split()).strip()

    try:
        with ZipFile(path) as zf:
            document_xml = zf.read("word/document.xml")
    except Exception:
        document_xml = b""
    if document_xml:
        try:
            root = ET.fromstring(document_xml)
            body = root.find("w:body", ns)
            if body is not None:
                lines: list[str] = []

                def _walk(children):
                    for child in children:
                        tag_name = _local_name(child.tag)
                        if tag_name == "p":
                            value = _xml_text(child)
                            if value:
                                lines.append(value)
                        elif tag_name == "tbl":
                            for row in child.findall("w:tr", ns):
                                cells = []
                                for cell in row.findall("w:tc", ns):
                                    value = _xml_text(cell)
                                    if value:
                                        cells.append(value)
                                if cells:
                                    lines.append(" | ".join(cells))
                        elif tag_name == "sdt":
                            sdt_content = child.find("w:sdtContent", ns)
                            if sdt_content is not None:
                                _walk(list(sdt_content))

                _walk(list(body))
                preview = "\n".join(lines).strip()
                if preview:
                    return preview
        except Exception:
            pass

    try:
        import docx
    except Exception:
        return ""
    try:
        document = docx.Document(path)
    except Exception:
        return ""
    lines: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines).strip()


def _pdf_preview_text(path: Path, *, full_document: bool = False) -> str:
    try:
        from pypdf import PdfReader
    except Exception:
        try:
            from PyPDF2 import PdfReader  # type: ignore
        except Exception:
            return ""
    try:
        reader = PdfReader(str(path))
    except Exception:
        return ""
    chunks: list[str] = []
    pages = reader.pages if full_document else reader.pages[:3]
    for page in pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            chunks.append(text.strip())
    return "\n".join(chunks).strip()


def _detect_template_schema_artifact(text: str) -> dict[str, Any]:
    if not text.strip():
        return {}
    lower = text.lower()
    placeholder_hits = len(_TEMPLATE_PLACEHOLDER_RE.findall(text))
    schema_hits = sum(
        1
        for token in (
            "prompt package json",
            "system_prompt",
            "user_prompt_template",
            "output_schema",
            '"type": "object"',
            '"additionalproperties": false',
            "return valid json only",
        )
        if token in lower
    )
    required_list_hits = 1 if '"required": [' in lower else 0
    if placeholder_hits >= 3 and (schema_hits + required_list_hits) >= 3:
        return {
            "template_schema_artifact": True,
            "template_schema_kind": "prompt_package",
            "template_schema_placeholder_hits": placeholder_hits,
            "template_schema_schema_hits": schema_hits + required_list_hits,
        }
    return {}


def _detect_meta_reference_artifact(text: str) -> dict[str, Any]:
    if not text.strip():
        return {}
    lower = text.lower()
    collapsed = re.sub(r"\s+", " ", lower)
    signal_hits = sum(1 for token in _META_REFERENCE_SIGNAL_TOKENS if token in lower or token in collapsed)
    architecture_hits = sum(1 for token in ("architecture", "parser", "evidence", "graph", "packet") if token in lower or token in collapsed)
    if signal_hits >= 5 and architecture_hits >= 3:
        return {
            "meta_reference_artifact": True,
            "meta_reference_kind": "parser_architecture",
            "meta_reference_signal_hits": signal_hits,
            "meta_reference_architecture_hits": architecture_hits,
        }
    return {}
