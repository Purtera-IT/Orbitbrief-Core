from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

from orbitbrief_core.compiler.packs.professional_services_text.compiler_runner import load_compiled_pack
from orbitbrief_core.runtime_spine.package_pipeline import run_package_pipeline
from orbitbrief_core.runtime_spine.pipeline import run_pipeline


def _compiled_pack():
    repo_root = Path(__file__).resolve().parents[2]
    return load_compiled_pack("professional_services_text", compiled_root=repo_root / "compiled_artifacts")


def _build_deal_kit(workbook_path: Path) -> None:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Deal Kit"
    summary.append(["Customer", "Musick, Peeler & Garrett"])
    summary.append(["QTY of Sites", 5])
    summary.append(["Project Duration (Months)", 12])
    summary.append(["Billing Type", "Fixed Fee - Monthly Billing"])

    roster = workbook.create_sheet("Site List")
    roster.append(["Site", "Job Description", "Billing Type", "Unit Sell Quantity", "Labor Rate Type"])
    roster.append([
        "Los Angeles HQ",
        "Dedicated Full-Time Onsite/Remote IT Support Resource - LABOR",
        "Fixed Fee - Monthly Billing",
        12,
        "Per Month",
    ])
    workbook.save(workbook_path)


def test_rich_sow_narrative_suppresses_intro_boilerplate_and_unstructured_quantities(tmp_path: Path) -> None:
    docx_path = tmp_path / "managed_services_sow.docx"
    document = Document()
    document.add_heading("INTRODUCTION", level=1)
    document.add_paragraph(
        'This Project Services Statement of Work ("SOW") is entered into by and between Musick, Peeler & Garrett (the "Customer") and PurTera LLC ("PurTera"). '
        'Under this SOW, PurTera shall provide one dedicated, full-time, English-speaking technical support resource to deliver onsite and remote end-user support across five (5) Customer locations.'
    )
    document.add_heading("PROJECT OVERVIEW", level=1)
    document.add_paragraph(
        "PurTera shall provide a dedicated full-time technical support resource for Musick, Peeler & Garrett to perform onsite services from the Los Angeles office and remote support for the broader five-location environment."
    )
    document.add_heading("DELIVERABLES", level=1)
    document.add_paragraph("Provision of one assigned full-time support resource for the duration of the engagement.", style="List Bullet")
    document.add_heading("CUSTOMER RESPONSIBILITIES", level=1)
    document.add_paragraph("Provide timely access to facilities, systems, and permissions necessary to perform the Services.", style="List Bullet")
    document.save(docx_path)

    result = run_pipeline(docx_path, compiled_pack=_compiled_pack(), include_runtime_result=True)
    field_claims = result["runtime_result"].postprocess_result["normalized_output"]["field_claims"]
    deliverables = [claim["candidate_value"] for claim in field_claims if claim["target_field_path"] == "deliverables[]"]
    summaries = [claim["candidate_value"] for claim in field_claims if claim["target_field_path"] in {"project_summary", "scope_overview"}]
    quantity_values = [claim["candidate_value"] for claim in field_claims if claim["target_field_path"] in {"scope_included[].quantity", "scope_included[].unit"}]

    assert deliverables == ["Provision of one assigned full-time support resource for the duration of the engagement"]
    assert summaries
    assert all("entered into by and between" not in str(value).lower() for value in summaries)
    assert all(len(str(value)) <= 40 for value in quantity_values)


def test_rich_sow_narrative_does_not_project_scope_headings_or_customer_responsibilities_as_scope(tmp_path: Path) -> None:
    docx_path = tmp_path / "managed_services_scope_guardrails.docx"
    document = Document()
    document.add_heading("PROJECT OVERVIEW", level=1)
    document.add_paragraph(
        "PurTera shall provide a dedicated full-time technical support resource for Musick, Peeler & Garrett to perform onsite services primarily from the Los Angeles office and to provide remote support, as required, for the broader five-location environment. The purpose of this engagement is to provide continuity of daily helpdesk support operations and general end-user technical issue resolution."
    )
    document.add_heading("SCOPE OF WORK", level=1)
    document.add_paragraph("In-scope services include:")
    document.add_paragraph("Incident intake and general troubleshooting.", style="List Bullet")
    document.add_paragraph("Out-of-scope services include:")
    document.add_paragraph("Application development.", style="List Bullet")
    document.add_heading("CUSTOMER RESPONSIBILITIES", level=1)
    document.add_paragraph("Provide timely access to systems, facilities, and permissions necessary to perform the Services.", style="List Bullet")
    document.save(docx_path)

    result = run_pipeline(docx_path, compiled_pack=_compiled_pack(), include_runtime_result=True)
    field_claims = result["runtime_result"].postprocess_result["normalized_output"]["field_claims"]
    scope_values = [claim["candidate_value"] for claim in field_claims if claim["target_field_path"] == "scope_included[]"]

    assert "In-scope services include" not in scope_values
    assert "Out-of-scope services include" not in scope_values
    assert not any("Provide timely access" in str(value) for value in scope_values)
    assert not any("The purpose of this engagement" in str(value) for value in scope_values)


def test_spreadsheet_customer_identity_row_does_not_emit_customer_responsibility_noise(tmp_path: Path) -> None:
    workbook_path = tmp_path / "deal_kit.xlsx"
    _build_deal_kit(workbook_path)

    result = run_pipeline(workbook_path, compiled_pack=_compiled_pack(), include_runtime_result=True)
    field_claims = result["runtime_result"].postprocess_result["normalized_output"]["field_claims"]
    noisy_paths = {"customer_documents_required[]", "customer_inputs_required[]", "customer_provided_materials[]", "customer_responsibilities[]"}
    noisy_values = [claim["candidate_value"] for claim in field_claims if claim["target_field_path"] in noisy_paths]

    assert noisy_values == []


def test_package_pipeline_joins_docx_and_spreadsheet_support(tmp_path: Path) -> None:
    workbook_path = tmp_path / "deal_kit.xlsx"
    _build_deal_kit(workbook_path)

    narrative_path = tmp_path / "project_overview.docx"
    document = Document()
    document.add_heading("PROJECT OVERVIEW", level=1)
    document.add_paragraph(
        "PurTera will provide dedicated onsite and remote support from the Los Angeles office across five customer locations for the initial managed services term."
    )
    document.save(narrative_path)

    package_result = run_package_pipeline([narrative_path, workbook_path], compiled_pack=_compiled_pack())
    site_claims = [claim for claim in package_result.joined_field_claims if claim["target_field_path"] == "site_locations[]"]
    site_count_claims = [claim for claim in package_result.joined_field_claims if claim["target_field_path"] == "site_count"]
    pricing_claims = [claim for claim in package_result.joined_field_claims if claim["target_field_path"] == "commercial_structure.pricing_model"]

    assert any(claim["candidate_value"] == "Los Angeles HQ" for claim in site_claims)
    joined_site = next(claim for claim in site_claims if claim["candidate_value"] == "Los Angeles HQ")
    assert joined_site["metadata"].get("package_joined") is True
    assert len(joined_site["evidence_span_ids"]) >= 2
    assert any(str(claim["candidate_value"]) == "5" for claim in site_count_claims)
    assert any(claim["candidate_value"] == "Fixed Fee - Monthly Billing" for claim in pricing_claims)


def test_package_pipeline_canonicalizes_narrative_pricing_from_spreadsheet(tmp_path: Path) -> None:
    workbook_path = tmp_path / "deal_kit.xlsx"
    _build_deal_kit(workbook_path)

    narrative_path = tmp_path / "pricing_terms.docx"
    document = Document()
    document.add_heading("PRICING & PAYMENT TERMS", level=1)
    document.add_paragraph(
        "Customer shall be invoiced monthly in arrears for the dedicated resource allocation established under this SOW."
    )
    document.save(narrative_path)

    package_result = run_package_pipeline([narrative_path, workbook_path], compiled_pack=_compiled_pack())
    pricing_claims = [claim for claim in package_result.joined_field_claims if claim["target_field_path"] == "commercial_structure.pricing_model"]

    assert any(claim["candidate_value"] == "Fixed Fee - Monthly Billing" for claim in pricing_claims)
    joined_pricing = next(
        claim for claim in pricing_claims
        if claim["metadata"].get("artifact_modality") == "docx" and claim["metadata"].get("package_joined") is True
    )
    assert joined_pricing["candidate_value"] == "Fixed Fee - Monthly Billing"
    assert joined_pricing["metadata"].get("package_join_reason") == "pricing_model_canonicalized_from_spreadsheet"
