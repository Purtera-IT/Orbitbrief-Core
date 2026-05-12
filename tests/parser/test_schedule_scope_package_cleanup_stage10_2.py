from __future__ import annotations

from pathlib import Path

from docx import Document
from openpyxl import Workbook

from orbitbrief_core.compiler.packs.professional_services_text.compiler_runner import load_compiled_pack
from orbitbrief_core.runtime_spine.extractors.packet_to_claims import PacketExtractionContext, extract_claims_from_packet
from orbitbrief_core.runtime_spine.package_pipeline import run_package_pipeline
from orbitbrief_core.runtime_spine.pipeline import run_pipeline


def _compiled_pack():
    repo_root = Path(__file__).resolve().parents[2]
    return load_compiled_pack("professional_services_text", compiled_root=repo_root / "compiled_artifacts")


def _build_deal_kit(workbook_path: Path) -> None:
    workbook = Workbook()
    summary = workbook.active
    summary.title = "Deal Kit"
    summary.append(["Customer", "Musick, Peeler & Garrett"])
    summary.append(["QTY of Sites", 5])
    summary.append(["Project Duration (Months)", 12])
    summary.append(["Billing Type", "Fixed Fee - Monthly Billing"])

    roster = workbook.create_sheet("Site List")
    roster.append(["Site", "Job Description", "Billing Type", "Unit Sell Quantity", "Labor Rate Type"])
    roster.append([
        "Los Angeles HQ",
        "Dedicated Full-Time Onsite/Remote IT Support Resource",
        "Fixed Fee - Monthly Billing",
        12,
        "Per Month",
    ])
    workbook.save(workbook_path)


def test_schedule_target_hints_do_not_leak_into_scope_or_assumptions() -> None:
    packet = {
        "packet_id": "packet:schedule_cluster:0001",
        "span_ids": ("span_schedule",),
        "primary_span_id": "span_schedule",
        "confidence": 0.83,
        "target_claim_family_names": ("scope_included_claim", "assumption_claim"),
        "evidence_rows": [
            {
                "span_id": "span_schedule",
                "text": "Initial term: One (1) year, unless otherwise agreed in writing.",
                "normalized_text": "initial term one 1 year unless otherwise agreed in writing",
                "section_path": ["ROOT", "TIMELINE & MILESTONES"],
                "parser_cues": ["schedule"],
                "packet_families": ["schedule_packet"],
                "authority_score": 0.9,
                "metadata": {"kind": "bullet"},
            }
        ],
        "metadata": {
            "packet_family": "schedule_packet",
            "uncertainty_markers": [],
            "packet_diagnostic": {"included": [{"span_id": "span_schedule"}]},
        },
    }

    claims, diagnostics = extract_claims_from_packet(
        packet,
        PacketExtractionContext(role_id="transcript_or_notes", modality="docx"),
    )

    claim_families = {claim.claim_family for claim in claims}
    assert claim_families == {"schedule_claim"}
    schedule_claim = claims[0]
    assert schedule_claim.metadata["schedule_semantic_class"] == "engagement_term"
    assert all(item.code not in {"semantic_family_override", "assumption_cross_family_suppressed"} for item in diagnostics)


def test_schedule_projection_only_keeps_true_commitments_in_completion_criteria(tmp_path: Path) -> None:
    docx_path = tmp_path / "timeline_cleanup.docx"
    document = Document()
    document.add_heading("PROJECT OVERVIEW", level=1)
    document.add_paragraph(
        "PurTera will provide dedicated onsite and remote helpdesk support for Musick, Peeler & Garrett."
    )
    document.add_heading("TIMELINE & MILESTONES", level=1)
    document.add_paragraph(
        "Planned service commencement: ASAP following SOW approval, resource onboarding, and completion of required access provisioning.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Initial transition period: approximately two (2) to four (4) weeks to review environment details and align support priorities.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Initial term: One (1) year, unless otherwise agreed in writing.",
        style="List Bullet",
    )
    document.add_paragraph(
        "Primary support coverage will align to the Customer's standard weekday operating schedule, with limited remote emergency support outside that schedule on an as-needed basis and subject to resource availability.",
        style="List Bullet",
    )
    document.add_heading("PRICING & PAYMENT TERMS", level=1)
    document.add_paragraph(
        "Customer shall be invoiced monthly in arrears for the dedicated resource allocation established under this SOW.",
        style="List Bullet",
    )
    document.save(docx_path)

    result = run_pipeline(docx_path, compiled_pack=_compiled_pack(), include_runtime_result=True)
    field_claims = result["runtime_result"].postprocess_result["normalized_output"]["field_claims"]
    completion_values = [claim["candidate_value"] for claim in field_claims if claim["target_field_path"] == "completion_criteria[]"]

    assert any("Initial transition period" in value for value in completion_values)
    assert len(completion_values) == 1
    assert all("Initial term" not in value for value in completion_values)
    assert all("weekday operating schedule" not in value for value in completion_values)
    assert all("invoiced monthly" not in value for value in completion_values)


def test_assumption_target_hints_are_limited_to_assumption_sections() -> None:
    packet = {
        "packet_id": "packet:intro_cluster:0001",
        "span_ids": ("span_intro",),
        "primary_span_id": "span_intro",
        "confidence": 0.8,
        "target_claim_family_names": ("assumption_claim",),
        "evidence_rows": [
            {
                "span_id": "span_intro",
                "text": 'This Project Services Statement of Work ("SOW") is entered into by and between Musick, Peeler & Garrett and PurTera LLC.',
                "normalized_text": "this project services statement of work sow is entered into by and between musick peeler and garrett and purtera llc",
                "section_path": ["ROOT", "INTRODUCTION"],
                "parser_cues": ["site_count"],
                "packet_families": ["site_packet"],
                "authority_score": 0.88,
                "metadata": {"kind": "paragraph"},
            }
        ],
        "metadata": {
            "packet_family": "site_packet",
            "uncertainty_markers": ["family_conflict"],
            "packet_diagnostic": {"included": [{"span_id": "span_intro"}]},
        },
    }

    claims, _diagnostics = extract_claims_from_packet(
        packet,
        PacketExtractionContext(role_id="transcript_or_notes", modality="docx"),
    )

    assert all(claim.claim_family != "assumption_claim" for claim in claims)


def test_package_pipeline_uses_relevant_scope_join_and_avoids_blanket_completion_support(tmp_path: Path) -> None:
    workbook_path = tmp_path / "deal_kit.xlsx"
    _build_deal_kit(workbook_path)

    narrative_path = tmp_path / "scope_and_timeline.docx"
    document = Document()
    document.add_heading("SCOPE OF WORK", level=1)
    document.add_paragraph(
        "Provide one dedicated onsite and remote IT support resource for the Los Angeles office.",
        style="List Bullet",
    )
    document.add_heading("TIMELINE & MILESTONES", level=1)
    document.add_paragraph(
        "Initial transition period: approximately two (2) to four (4) weeks to review environment details and align support priorities.",
        style="List Bullet",
    )
    document.save(narrative_path)

    package_result = run_package_pipeline([narrative_path, workbook_path], compiled_pack=_compiled_pack())

    scope_claim = next(
        claim for claim in package_result.joined_field_claims
        if claim["target_field_path"] == "scope_included[]"
        and claim["metadata"].get("artifact_modality") == "docx"
        and "dedicated onsite and remote it support resource" in str(claim["candidate_value"]).lower()
    )
    assert scope_claim["metadata"].get("package_joined") is True
    assert scope_claim["metadata"].get("package_join_reason") == "scope_item_supported_by_spreadsheet"

    completion_claims = [
        claim for claim in package_result.joined_field_claims
        if claim["target_field_path"] == "completion_criteria[]"
        and claim["metadata"].get("artifact_modality") == "docx"
    ]
    assert completion_claims
    assert all(claim["metadata"].get("package_joined") is not True for claim in completion_claims)
